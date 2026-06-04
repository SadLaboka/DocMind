from collections.abc import Callable
from io import BytesIO
from pathlib import Path
from typing import Final, TypeAlias

from src.core.exceptions import ExtractionError
from src.schemas.documents import MimeType

ExtractorFunc: TypeAlias = Callable[[BytesIO], str]


class TextExtractor:
    def __init__(self):
        self.extractors: Final[dict[MimeType, ExtractorFunc]] = {
            MimeType.txt: self._extract_txt,
            MimeType.docx: self._extract_docx,
            MimeType.xlsx: self._extract_xlsx,
            MimeType.pdf: self._extract_pdf,
        }

    def extract(self, temp_filename: Path | str, mime_type: MimeType) -> str:
        """Extracts the text of the uploaded file"""
        extractor = self.extractors.get(mime_type)
        if not extractor:
            raise ExtractionError(
                error_code="invalid_mime_type",
                log_context={"mime_type": mime_type},
            )
        file = None
        try:
            file = self._get_file_data(temp_filename)
            file.seek(0)
            text = extractor(file)
        finally:
            if file is not None:
                file.close()
        return text

    @staticmethod
    def _get_file_data(temp_filename: Path | str) -> BytesIO:
        """Reads the uploaded file and returns the data"""
        path = Path(temp_filename)
        if not path.is_file():
            raise ExtractionError(
                error_code="file_not_found",
                log_context={"path": path},
            )
        return BytesIO(path.read_bytes())

    def _extract_txt(self, file: BytesIO) -> str:
        """Extracts the text from bytesio if file has txt mimetype"""
        return file.getvalue().decode("utf-8", errors="replace")

    def _extract_docx(self, file: BytesIO) -> str:
        """Extracts the text from bytesio if file has docx mimetype"""
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError

        try:
            document = Document(file)
        except PackageNotFoundError as err:
            raise ExtractionError(
                error_code="invalid_file",
                log_context={"detail": str(err)},
            )
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        if document.tables:
            for table_num, table in enumerate(document.tables):
                table_parts = [f"--- Table {table_num} ---"]
                for row_num, row in enumerate(table.rows):
                    row_parts = []
                    for cell_num, cell in enumerate(row.cells):
                        cell_text = "\n".join(p.text for p in cell.paragraphs).strip()
                        if cell_text:
                            row_parts.append(f"[Cell {cell_num}]: {cell_text}")

                    if row_parts:
                        table_parts.append(" | ".join(row_parts))

                text += "\n" + "\n".join(table_parts)

        return text

    def _extract_xlsx(self, file: BytesIO) -> str:
        """Extracts the text from bytesio if file has xlsx mimetype"""
        from openpyxl import load_workbook

        text_parts = []
        try:
            workbook = load_workbook(filename=file, read_only=True, data_only=True)

            for sheet in workbook.worksheets:
                text_parts.append(f"--- Sheet: {sheet.title} ---")

                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None and str(cell).strip() for cell in row):
                        row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                        text_parts.append(row_text)

        except Exception as err:
            raise ExtractionError(
                error_code="invalid_file",
                log_context={"detail": str(err)},
            )

        return "\n".join(text_parts)

    def _extract_pdf(self, file: BytesIO) -> str:
        """Extracts the text from bytesio if file has PDF mimetype"""
        import pdfplumber
        from pdfminer.pdfparser import PDFSyntaxError

        text_parts = []
        try:
            with pdfplumber.open(file) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()

                    if page_text:
                        text_parts.append(f"--- Page {page_num} ---")
                        text_parts.append(page_text.strip())

                    tables = page.extract_tables()
                    if tables:

                        for table_num, table in enumerate(tables):
                            text_parts.append(f"--- Page {page_num}, Table {table_num} ---")

                            for row_num, row in enumerate(table):
                                text_parts.append(f"--- Row {row_num} ---")
                                if any(cell is not None and str(cell).strip() for cell in row):
                                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                                    text_parts.append(row_text)
        except (PDFSyntaxError, Exception) as err:
            raise ExtractionError(
                error_code="invalid_file",
                log_context={"detail": str(err)},
            )

        return "\n".join(text_parts)
